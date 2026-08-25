#!/usr/bin/env python3
"""Export MSSP Master Blueprint Markdown to PDF and DOCX.

Reads:  docs/MSSP_PLATFORM_MASTER_BLUEPRINT.md
Writes: docs/MSSP_PLATFORM_MASTER_BLUEPRINT.pdf
        docs/MSSP_PLATFORM_MASTER_BLUEPRINT.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "MSSP_PLATFORM_MASTER_BLUEPRINT.md"
OUT_DIRS = [ROOT / "docs"]


def parse_blocks(text: str) -> list[tuple[str, str]]:
    """Return list of (kind, content) where kind in heading1-3, para, bullet, table, hr, code."""
    lines = text.splitlines()
    blocks: list[tuple[str, str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            blocks.append(("hr", ""))
            i += 1
            continue
        if line.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # closing fence
            blocks.append(("code", "\n".join(buf)))
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s\-:]+$", "".join(row).replace("|", "")):
                    # skip separator visually handled by skipping dash-only rows
                    if not all(re.match(r"^:?-+:?$", c) for c in row):
                        rows.append(row)
                i += 1
            blocks.append(("table", "\n".join("\t".join(r) for r in rows)))
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            blocks.append((f"h{level}", m.group(2).strip()))
            i += 1
            continue
        if re.match(r"^[-*]\s+", line):
            buf = [re.sub(r"^[-*]\s+", "", line)]
            i += 1
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                buf.append(re.sub(r"^[-*]\s+", "", lines[i]))
                i += 1
            blocks.append(("bullets", "\n".join(buf)))
            continue
        if not line.strip():
            i += 1
            continue
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("|") and not lines[i].startswith("```") and not lines[i].strip() == "---" and not re.match(r"^[-*]\s+", lines[i]):
            buf.append(lines[i])
            i += 1
        blocks.append(("para", " ".join(x.strip() for x in buf)))
    return blocks


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    return s


def export_docx(blocks: list[tuple[str, str]], path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for kind, content in blocks:
        if kind == "h1":
            p = doc.add_heading(strip_md_inline(content), level=0)
        elif kind == "h2":
            doc.add_heading(strip_md_inline(content), level=1)
        elif kind == "h3":
            doc.add_heading(strip_md_inline(content), level=2)
        elif kind == "hr":
            doc.add_paragraph("—" * 40)
        elif kind == "para":
            doc.add_paragraph(strip_md_inline(content))
        elif kind == "bullets":
            for item in content.split("\n"):
                doc.add_paragraph(strip_md_inline(item), style="List Bullet")
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        elif kind == "table":
            rows = [r.split("\t") for r in content.split("\n") if r.strip()]
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(cols):
                    cell = table.rows[ri].cells[ci]
                    cell.text = strip_md_inline(row[ci] if ci < len(row) else "")
            doc.add_paragraph("")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def export_pdf(blocks: list[tuple[str, str]], path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Preformatted,
        Table,
        TableStyle,
        PageBreak,
        HRFlowable,
        KeepTogether,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="BodyBlue",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Blue",
            parent=styles["Heading1"],
            fontSize=16,
            spaceBefore=14,
            spaceAfter=8,
            textColor=colors.HexColor("#0f2744"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2Blue",
            parent=styles["Heading2"],
            fontSize=13,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor("#163a5f"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="H3Blue",
            parent=styles["Heading3"],
            fontSize=11,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor("#1c4b78"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeBlue",
            fontName="Courier",
            fontSize=7.5,
            leading=9.5,
            backColor=colors.HexColor("#f4f6f8"),
            spaceBefore=4,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletBlue",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            leftIndent=14,
            bulletIndent=0,
            spaceAfter=2,
        )
    )

    def esc(s: str) -> str:
        s = strip_md_inline(s)
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    story = []
    for kind, content in blocks:
        if kind == "h1":
            story.append(Paragraph(esc(content), styles["H1Blue"]))
        elif kind == "h2":
            story.append(Paragraph(esc(content), styles["H2Blue"]))
        elif kind == "h3":
            story.append(Paragraph(esc(content), styles["H3Blue"]))
        elif kind == "hr":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.6, color=colors.grey))
            story.append(Spacer(1, 6))
        elif kind == "para":
            story.append(Paragraph(esc(content), styles["BodyBlue"]))
        elif kind == "bullets":
            for item in content.split("\n"):
                story.append(Paragraph("• " + esc(item), styles["BulletBlue"]))
        elif kind == "code":
            story.append(Preformatted(content, styles["CodeBlue"]))
        elif kind == "table":
            rows_raw = [r.split("\t") for r in content.split("\n") if r.strip()]
            if not rows_raw:
                continue
            cols = max(len(r) for r in rows_raw)
            data = []
            for ri, row in enumerate(rows_raw):
                cells = []
                for ci in range(cols):
                    txt = esc(row[ci] if ci < len(row) else "")
                    style = styles["BodyBlue"]
                    cells.append(Paragraph(txt, style))
                data.append(cells)
            t = Table(data, repeatRows=1, hAlign="LEFT")
            t.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9aa8b8")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 8))

    path.parent.mkdir(parents=True, exist_ok=True)

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.grey)
        canvas.drawString(inch * 0.9, 0.5 * inch, "Kevantic Cyber Security — MSSP Platform Master Blueprint")
        canvas.drawRightString(A4[0] - inch * 0.9, 0.5 * inch, f"Page {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="MSSP Platform Master Blueprint",
        author="Kestrel Cyber MSSP Control Plane",
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing markdown: {MD_PATH}", file=sys.stderr)
        return 1
    text = MD_PATH.read_text(encoding="utf-8")
    blocks = parse_blocks(text)
    print(f"Parsed {len(blocks)} blocks from {MD_PATH}")

    out_dir = ROOT / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf = out_dir / "MSSP_PLATFORM_MASTER_BLUEPRINT.pdf"
    docx = out_dir / "MSSP_PLATFORM_MASTER_BLUEPRINT.docx"
    export_docx(blocks, docx)
    print(f"Wrote {docx} ({docx.stat().st_size} bytes)")
    export_pdf(blocks, pdf)
    print(f"Wrote {pdf} ({pdf.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
